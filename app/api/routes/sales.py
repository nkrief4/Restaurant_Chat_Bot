"""Sales analytics endpoints enabling file analysis and confirmations."""

from __future__ import annotations

import asyncio
import io
import logging
import math
import re
from collections import defaultdict
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Literal, Optional, Sequence, Tuple
from uuid import UUID, uuid4

import pandas as pd
import pdfplumber
from docx import Document
from fastapi import APIRouter, Depends, File, Header, HTTPException, UploadFile
from httpx import HTTPError as HttpxError
from pydantic import BaseModel, Field
from rapidfuzz import fuzz, process
from postgrest import APIError as PostgrestAPIError

from app.config.supabase_client import SUPABASE_SERVICE_ROLE_KEY
from app.services.postgrest_client import (
    create_postgrest_client,
    extract_bearer_token,
    raise_postgrest_error,
)

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/sales", tags=["Sales"])


class SalesSuggestion(BaseModel):
    menu_item_id: UUID
    menu_item_name: str
    confidence: float


class SalesAnalysisLine(BaseModel):
    line_id: UUID
    raw_name: str
    quantity: float
    served_at: Optional[datetime] = None
    menu_item_id: Optional[UUID] = None
    menu_item_name: Optional[str] = None
    confidence: Optional[float] = None
    status: Literal["recognized", "unmatched"]
    suggestions: List[SalesSuggestion] = Field(default_factory=list)


class SalesAnalysisResponse(BaseModel):
    file_name: str
    total_rows: int
    recognized_count: int
    unmatched_count: int
    columns_detected: Dict[str, Optional[str]]
    recognized: List[SalesAnalysisLine]
    unmatched: List[SalesAnalysisLine]


class SalesConfirmLine(BaseModel):
    line_id: UUID
    raw_name: Optional[str] = None
    menu_item_id: UUID
    quantity: float = Field(..., gt=0)
    served_at: Optional[datetime] = None


class SalesConfirmRequest(BaseModel):
    lines: List[SalesConfirmLine]
    source_label: Optional[str] = Field(default=None, max_length=200)


class SalesTrendPoint(BaseModel):
    date_iso: date
    label: str
    quantity: float


class SalesLeaderboardRow(BaseModel):
    menu_item_id: UUID
    menu_item_name: str
    quantity: float
    previous_quantity: float
    delta: float


class SalesInsightsResponse(BaseModel):
    range_days: int = 7
    generated_at: datetime
    weekly_total: float
    trend: List[SalesTrendPoint]
    top_items: List[SalesLeaderboardRow]
    table: List[SalesLeaderboardRow]


async def get_current_restaurant_id(
    x_restaurant_id: Optional[str] = Header(default=None, alias="X-Restaurant-Id"),
) -> UUID:
    """Resolve the restaurant identifier from the current request."""

    if not x_restaurant_id:
        raise HTTPException(status_code=401, detail="Restaurant non authentifié.")
    try:
        return UUID(x_restaurant_id)
    except (TypeError, ValueError) as exc:  # pragma: no cover - defensive
        raise HTTPException(status_code=400, detail="Identifiant restaurant invalide.") from exc


async def get_access_token(
    authorization: Optional[str] = Header(default=None, alias="Authorization"),
) -> str:
    """Extract the Supabase bearer token from the Authorization header."""

    return extract_bearer_token(authorization)


async def ensure_restaurant_access(restaurant_id: UUID, access_token: str) -> None:
    """Ensure the caller can access the requested restaurant before using service credentials."""

    def _request() -> bool:
        with create_postgrest_client(access_token) as client:
            response = (
                client.table("restaurants")
                .select("id")
                .eq("id", str(restaurant_id))
                .limit(1)
                .execute()
            )
            return bool(response.data)

    try:
        has_access = await asyncio.to_thread(_request)
    except PostgrestAPIError as exc:  # pragma: no cover
        raise_postgrest_error(exc, context="restaurant access check")
    except HttpxError as exc:  # pragma: no cover
        raise HTTPException(status_code=503, detail="Supabase est temporairement inaccessible.") from exc

    if not has_access:
        raise HTTPException(status_code=403, detail="Accès refusé à ce restaurant.")


@router.post("/analyze", response_model=SalesAnalysisResponse)
async def analyze_sales_file(
    file: UploadFile = File(...),
    restaurant_id: UUID = Depends(get_current_restaurant_id),
    access_token: str = Depends(get_access_token),
) -> SalesAnalysisResponse:
    await ensure_restaurant_access(restaurant_id, access_token)
    service = SalesService(restaurant_id, access_token)
    return await service.analyze(file)


@router.post("/confirm", response_model=SalesInsightsResponse)
async def confirm_sales_import(
    payload: SalesConfirmRequest,
    restaurant_id: UUID = Depends(get_current_restaurant_id),
    access_token: str = Depends(get_access_token),
) -> SalesInsightsResponse:
    await ensure_restaurant_access(restaurant_id, access_token)
    service = SalesService(restaurant_id, access_token)
    return await service.confirm(payload)


@router.get("/insights", response_model=SalesInsightsResponse)
async def fetch_sales_insights(
    restaurant_id: UUID = Depends(get_current_restaurant_id),
    access_token: str = Depends(get_access_token),
) -> SalesInsightsResponse:
    await ensure_restaurant_access(restaurant_id, access_token)
    service = SalesService(restaurant_id, access_token)
    return await service.get_insights()


class ParsedSaleRow(BaseModel):
    line_id: UUID
    raw_name: str
    quantity: float
    served_at: Optional[datetime] = None


class SalesService:
    """Encapsulates file parsing, fuzzy matching and persistence logic."""

    MATCH_THRESHOLD = 78

    def __init__(self, restaurant_id: UUID, access_token: str) -> None:
        self.restaurant_id = str(restaurant_id)
        self.access_token = access_token
        if SUPABASE_SERVICE_ROLE_KEY:
            self.db_token = SUPABASE_SERVICE_ROLE_KEY
            self.api_key = SUPABASE_SERVICE_ROLE_KEY
        else:
            self.db_token = access_token
            self.api_key = None

    async def analyze(self, upload: UploadFile) -> SalesAnalysisResponse:
        payload = await upload.read()
        if not payload:
            raise HTTPException(status_code=400, detail="Le fichier est vide.")

        parser = SalesFileParser(upload.filename or "import", upload.content_type, payload)
        try:
            rows, columns_detected = parser.parse()
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

        if not rows:
            raise HTTPException(status_code=400, detail="Aucune ligne de vente détectée.")

        menu_items = await self._fetch_menu_items()
        matcher = SalesMatcher(menu_items, threshold=self.MATCH_THRESHOLD)
        recognized, unmatched = matcher.match_rows(rows)

        return SalesAnalysisResponse(
            file_name=upload.filename or "import",
            total_rows=len(rows),
            recognized_count=len(recognized),
            unmatched_count=len(unmatched),
            columns_detected=columns_detected,
            recognized=recognized,
            unmatched=unmatched,
        )

    async def confirm(self, payload: SalesConfirmRequest) -> SalesInsightsResponse:
        if not payload.lines:
            raise HTTPException(status_code=400, detail="Ajoutez au moins une ligne à confirmer.")

        def _request() -> None:
            body: List[Dict[str, Any]] = []
            for line in payload.lines:
                ordered_at = self._format_datetime(line.served_at)
                body.append(
                    {
                        "restaurant_id": self.restaurant_id,
                        "menu_item_id": str(line.menu_item_id),
                        "quantity": int(math.ceil(line.quantity)),
                        "ordered_at": ordered_at,
                        "source": payload.source_label or "sales_import_dashboard",
                    }
                )
            with create_postgrest_client(
                self.db_token,
                prefer="return=representation",
                api_key=self.api_key,
            ) as client:
                client.table("orders").insert(body).execute()

        try:
            await asyncio.to_thread(_request)
        except PostgrestAPIError as exc:  # pragma: no cover - HTTP interaction
            raise_postgrest_error(exc, context="confirm sales import")
        except HttpxError as exc:  # pragma: no cover
            raise HTTPException(status_code=503, detail="Supabase est temporairement inaccessible.") from exc

        return await self.get_insights()

    async def get_insights(self) -> SalesInsightsResponse:
        range_days = 7
        now = datetime.now(timezone.utc)
        start_current = (now - timedelta(days=range_days - 1)).date()
        start_previous = start_current - timedelta(days=range_days)

        def _request() -> List[Dict[str, Any]]:
            with create_postgrest_client(self.db_token, api_key=self.api_key) as client:
                response = (
                    client.table("orders")
                    .select("menu_item_id,quantity,ordered_at,menu_items(name)")
                    .eq("restaurant_id", self.restaurant_id)
                    .gte("ordered_at", f"{start_previous.isoformat()}T00:00:00Z")
                    .lte("ordered_at", now.isoformat())
                    .execute()
                )
                return response.data or []

        try:
            records = await asyncio.to_thread(_request)
        except PostgrestAPIError as exc:  # pragma: no cover
            raise_postgrest_error(exc, context="fetch sales insights")
        except HttpxError as exc:  # pragma: no cover
            raise HTTPException(status_code=503, detail="Supabase est temporairement inaccessible.") from exc

        current_totals: Dict[str, float] = defaultdict(float)
        previous_totals: Dict[str, float] = defaultdict(float)
        timeline: Dict[date, float] = {start_current + timedelta(days=i): 0.0 for i in range(range_days)}
        names: Dict[str, str] = {}

        for record in records:
            menu_item_id = record.get("menu_item_id")
            if not menu_item_id:
                continue
            quantity = _safe_float(record.get("quantity"), default=0.0)
            if quantity <= 0:
                continue
            ordered_at = _parse_datetime(record.get("ordered_at"))
            if not ordered_at:
                continue
            ordered_date = ordered_at.date()
            names[str(menu_item_id)] = (
                ((record.get("menu_items") or {}).get("name")) or "Plat"
            )
            if ordered_date in timeline:
                timeline[ordered_date] += quantity
                current_totals[str(menu_item_id)] += quantity
            elif start_previous <= ordered_date < start_current:
                previous_totals[str(menu_item_id)] += quantity

        trend_points = [
            SalesTrendPoint(
                date_iso=day,
                label=day.strftime("%d %b"),
                quantity=round(timeline.get(day, 0.0), 2),
            )
            for day in (start_current + timedelta(days=i) for i in range(range_days))
        ]

        leaderboard = _build_leaderboard(current_totals, previous_totals, names)
        weekly_total = float(sum(point.quantity for point in trend_points))

        return SalesInsightsResponse(
            range_days=range_days,
            generated_at=now,
            weekly_total=weekly_total,
            trend=trend_points,
            top_items=leaderboard[:5],
            table=leaderboard[:8],
        )

    async def _fetch_menu_items(self) -> List[Dict[str, Any]]:
        def _request() -> List[Dict[str, Any]]:
            with create_postgrest_client(self.db_token, api_key=self.api_key) as client:
                response = (
                    client.table("menu_items")
                    .select("id,name,is_active")
                    .eq("restaurant_id", self.restaurant_id)
                    .execute()
                )
                return response.data or []

        try:
            return await asyncio.to_thread(_request)
        except PostgrestAPIError as exc:  # pragma: no cover
            raise_postgrest_error(exc, context="fetch menu items for sales")
        except HttpxError as exc:  # pragma: no cover
            raise HTTPException(status_code=503, detail="Supabase est temporairement inaccessible.") from exc

    @staticmethod
    def _format_datetime(value: Optional[datetime]) -> str:
        if value is None:
            return datetime.now(timezone.utc).isoformat()
        if value.tzinfo is None:
            value = value.replace(tzinfo=timezone.utc)
        return value.astimezone(timezone.utc).isoformat()


class SalesMatcher:
    """Perform fuzzy matching between parsed rows and menu items."""

    def __init__(self, menu_items: Sequence[Dict[str, Any]], threshold: int = 78) -> None:
        self.threshold = threshold
        self.menu_lookup: Dict[str, Dict[str, Any]] = {}
        choices: Dict[str, str] = {}
        for row in menu_items:
            identifier = row.get("id")
            if not identifier:
                continue
            name = row.get("display_name") or row.get("name") or "Plat"
            key = str(identifier)
            self.menu_lookup[key] = {"id": identifier, "name": name}
            choices[key] = name
        self.choices = choices

    def match_rows(self, rows: Sequence[ParsedSaleRow]) -> Tuple[List[SalesAnalysisLine], List[SalesAnalysisLine]]:
        recognized: List[SalesAnalysisLine] = []
        unmatched: List[SalesAnalysisLine] = []
        for row in rows:
            match = self._match(row.raw_name)
            suggestions = self._suggest(row.raw_name)
            if match and match[1] >= self.threshold:
                menu_item = self.menu_lookup.get(match[2])
                recognized.append(
                    SalesAnalysisLine(
                        line_id=row.line_id,
                        raw_name=row.raw_name,
                        quantity=row.quantity,
                        served_at=row.served_at,
                        menu_item_id=UUID(str(menu_item["id"])),
                        menu_item_name=menu_item["name"],
                        confidence=float(match[1]),
                        status="recognized",
                        suggestions=suggestions,
                    )
                )
            else:
                best = suggestions[0] if suggestions else None
                unmatched.append(
                    SalesAnalysisLine(
                        line_id=row.line_id,
                        raw_name=row.raw_name,
                        quantity=row.quantity,
                        served_at=row.served_at,
                        menu_item_id=best.menu_item_id if best else None,
                        menu_item_name=best.menu_item_name if best else None,
                        confidence=best.confidence if best else None,
                        status="unmatched",
                        suggestions=suggestions,
                    )
                )
        return recognized, unmatched

    def _match(self, query: str) -> Optional[Tuple[str, float, str]]:
        if not query.strip():
            return None
        return process.extractOne(
            query,
            self.choices,
            scorer=fuzz.WRatio,
        )

    def _suggest(self, query: str) -> List[SalesSuggestion]:
        if not query.strip():
            return []
        suggestions: List[SalesSuggestion] = []
        matches = process.extract(
            query,
            self.choices,
            scorer=fuzz.WRatio,
            limit=3,
        )
        for match in matches:
            name, score, identifier = match
            if score < 45:
                continue
            suggestions.append(
                SalesSuggestion(
                    menu_item_id=UUID(str(self.menu_lookup[identifier]["id"])),
                    menu_item_name=name,
                    confidence=float(score),
                )
            )
        return suggestions


class SalesFileParser:
    """Normalize supported upload formats into unified sales rows."""

    NAME_ALIASES = ("plat", "dish", "item", "article", "produit", "menu", "recette", "name", "designation")
    QUANTITY_ALIASES = ("qty", "qte", "quantite", "quantity", "ventes", "sales", "units", "nb", "nombre")
    DATE_ALIASES = ("date", "jour", "day", "sold", "commande", "order", "service")

    def __init__(self, filename: str, content_type: Optional[str], payload: bytes) -> None:
        self.filename = filename
        self.content_type = (content_type or "").lower()
        self.payload = payload

    def parse(self) -> Tuple[List[ParsedSaleRow], Dict[str, Optional[str]]]:
        suffix = Path(self.filename).suffix.lower()
        if suffix in (".xlsx", ".xls"):
            return self._parse_excel()
        if suffix == ".csv" or "csv" in self.content_type:
            return self._parse_csv()
        if suffix == ".docx" or "word" in self.content_type:
            return self._parse_docx()
        if suffix == ".pdf" or "pdf" in self.content_type:
            return self._parse_pdf()
        # fallback to try csv heuristics
        return self._parse_csv()

    def _parse_csv(self) -> Tuple[List[ParsedSaleRow], Dict[str, Optional[str]]]:
        buffer = io.BytesIO(self.payload)
        try:
            df = pd.read_csv(buffer)
        except Exception as exc:  # pragma: no cover - delegated to pandas
            raise ValueError("Impossible de lire ce fichier CSV.") from exc
        return self._rows_from_dataframe(df)

    def _parse_excel(self) -> Tuple[List[ParsedSaleRow], Dict[str, Optional[str]]]:
        buffer = io.BytesIO(self.payload)
        suffix = Path(self.filename or "").suffix.lower()
        try:
            engine = "openpyxl" if suffix != ".xls" else None
            df = pd.read_excel(buffer, engine=engine)
        except ImportError as exc:  # pragma: no cover - openpyxl missing
            raise ValueError("Le support des fichiers Excel nécessite le paquet openpyxl.") from exc
        except ValueError as exc:
            if suffix == ".xls":
                raise ValueError("Les fichiers .xls hérités ne sont pas supportés. Exportez en .xlsx ou .csv.") from exc
            raise ValueError("Impossible de lire ce fichier Excel.") from exc
        except Exception as exc:  # pragma: no cover
            raise ValueError("Impossible de lire ce fichier Excel.") from exc
        return self._rows_from_dataframe(df)

    def _parse_docx(self) -> Tuple[List[ParsedSaleRow], Dict[str, Optional[str]]]:
        document = Document(io.BytesIO(self.payload))
        rows: List[ParsedSaleRow] = []
        columns: Dict[str, Optional[str]] = {}

        table_records: List[Dict[str, str]] = []
        for table in document.tables:
            raw_rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not raw_rows:
                continue
            header = raw_rows[0]
            if not any(header):
                continue
            keys = self._hydrate_headers(header)
            for table_row in raw_rows[1:]:
                record = {keys[idx]: cell for idx, cell in enumerate(table_row) if idx < len(keys)}
                table_records.append(record)

        if table_records:
            df = pd.DataFrame.from_records(table_records)
            parsed, columns = self._rows_from_dataframe(df)
            if parsed:
                return parsed, columns

        text_lines = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
        return self._rows_from_text(text_lines)

    def _parse_pdf(self) -> Tuple[List[ParsedSaleRow], Dict[str, Optional[str]]]:
        text_lines: List[str] = []
        table_records: List[Dict[str, str]] = []

        with pdfplumber.open(io.BytesIO(self.payload)) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables() or []
                for table in tables:
                    if len(table) < 2:
                        continue
                    header = table[0]
                    keys = self._hydrate_headers(header)
                    for raw in table[1:]:
                        record = {keys[idx]: cell for idx, cell in enumerate(raw) if idx < len(keys)}
                        table_records.append(record)
                content = page.extract_text() or ""
                text_lines.extend([line.strip() for line in content.splitlines() if line.strip()])

        if table_records:
            df = pd.DataFrame.from_records(table_records)
            parsed, columns = self._rows_from_dataframe(df)
            if parsed:
                return parsed, columns

        return self._rows_from_text(text_lines)

    def _rows_from_dataframe(self, df: pd.DataFrame) -> Tuple[List[ParsedSaleRow], Dict[str, Optional[str]]]:
        if df.empty:
            return [], {"name": None, "quantity": None, "date": None}
        normalized = {self._normalize_header(col): col for col in df.columns}
        name_col = self._resolve_column(normalized, self.NAME_ALIASES) or df.columns[0]
        qty_col = self._resolve_column(normalized, self.QUANTITY_ALIASES)
        if qty_col is None:
            numeric_cols = [col for col in df.columns if pd.api.types.is_numeric_dtype(df[col])]
            qty_col = numeric_cols[0] if numeric_cols else df.columns[1] if len(df.columns) > 1 else df.columns[0]
        date_col = self._resolve_column(normalized, self.DATE_ALIASES)

        detected = {
            "name": name_col,
            "quantity": qty_col,
            "date": date_col,
        }

        rows: List[ParsedSaleRow] = []
        for _, record in df.iterrows():
            raw_name = str(record.get(name_col, "")).strip()
            if not raw_name or raw_name.lower() in {"nan", "none"}:
                continue
            quantity = _safe_float(record.get(qty_col), default=1.0)
            if quantity <= 0:
                continue
            served_at = _parse_datetime(record.get(date_col)) if date_col else None
            rows.append(
                ParsedSaleRow(
                    line_id=uuid4(),
                    raw_name=raw_name,
                    quantity=round(quantity, 3),
                    served_at=served_at,
                )
            )
            if len(rows) >= 500:
                break
        return rows, detected

    def _rows_from_text(self, lines: Iterable[str]) -> Tuple[List[ParsedSaleRow], Dict[str, Optional[str]]]:
        pattern = re.compile(
            r"^(?P<name>[\wÀ-ÖØ-öø-ÿ'&().\-\s]+?)\s{1,}[-x*:]?\s*(?P<qty>\d+(?:[.,]\d+)?)",
            re.IGNORECASE,
        )
        date_pattern = re.compile(r"(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})")
        rows: List[ParsedSaleRow] = []
        for entry in lines:
            match = pattern.search(entry)
            if not match:
                continue
            raw_name = match.group("name").strip(" -:\t")
            quantity = _safe_float(match.group("qty"), default=0.0)
            if not raw_name or quantity <= 0:
                continue
            served_at = None
            date_match = date_pattern.search(entry)
            if date_match:
                served_at = _parse_datetime(date_match.group(1))
            rows.append(
                ParsedSaleRow(
                    line_id=uuid4(),
                    raw_name=raw_name,
                    quantity=round(quantity, 3),
                    served_at=served_at,
                )
            )
            if len(rows) >= 500:
                break
        detected = {"name": "text_line", "quantity": "pattern", "date": "inline_date" if rows else None}
        return rows, detected

    @staticmethod
    def _normalize_header(value: str) -> str:
        return re.sub(r"[^a-z0-9]", "", str(value).strip().lower())

    def _resolve_column(self, mapping: Dict[str, str], aliases: Sequence[str]) -> Optional[str]:
        for alias in aliases:
            candidate = mapping.get(alias)
            if candidate:
                return candidate
        for key, column in mapping.items():
            for alias in aliases:
                if alias in key:
                    return column
        return None

    @staticmethod
    def _hydrate_headers(headers: Sequence[str]) -> List[str]:
        hydrated = []
        for idx, header in enumerate(headers):
            name = str(header or f"col_{idx}").strip()
            hydrated.append(name or f"col_{idx}")
        return hydrated


def _safe_float(value: Any, *, default: float) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _parse_datetime(value: Any) -> Optional[datetime]:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return None
    if isinstance(value, datetime):
        return value if value.tzinfo else value.replace(tzinfo=timezone.utc)
    if isinstance(value, pd.Timestamp):
        return value.to_pydatetime()
    if isinstance(value, date):
        return datetime.combine(value, datetime.min.time(), tzinfo=timezone.utc)
    if isinstance(value, str):
        value = value.strip()
        if not value:
            return None
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y/%m/%d"):
            try:
                parsed = datetime.strptime(value, fmt)
                return parsed.replace(tzinfo=timezone.utc)
            except ValueError:
                continue
        try:
            parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
            if parsed.tzinfo is None:
                parsed = parsed.replace(tzinfo=timezone.utc)
            return parsed
        except ValueError:
            return None
    return None


def _build_leaderboard(
    current_totals: Dict[str, float],
    previous_totals: Dict[str, float],
    names: Dict[str, str],
) -> List[SalesLeaderboardRow]:
    leaderboard: List[SalesLeaderboardRow] = []
    for menu_id, quantity in current_totals.items():
        prev = previous_totals.get(menu_id, 0.0)
        delta = _compute_delta(prev, quantity)
        try:
            identifier = UUID(menu_id)
        except ValueError:
            continue
        leaderboard.append(
            SalesLeaderboardRow(
                menu_item_id=identifier,
                menu_item_name=names.get(menu_id, "Plat"),
                quantity=round(quantity, 2),
                previous_quantity=round(prev, 2),
                delta=delta,
            )
        )
    leaderboard.sort(key=lambda entry: entry.quantity, reverse=True)
    return leaderboard


def _compute_delta(previous: float, current: float) -> float:
    if previous <= 0:
        return 100.0 if current > 0 else 0.0
    return round(((current - previous) / previous) * 100, 2)


__all__ = [
    "router",
]
